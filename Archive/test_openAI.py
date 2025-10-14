import os
import re
import pandas as pd
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

load_dotenv()
key = os.getenv("OPENAI_API_KEY")

if not key:
    raise SystemExit("OPENAI_API_KEY not found in .env")

client = OpenAI(api_key=key)

def load_csv_data():
    if not os.path.exists("inputted_ttps.csv"):
        raise SystemExit("inputted_ttps.csv not found. Run matching_test.py first.")
    if not os.path.exists("matched_groups.csv"):
        raise SystemExit("matched_groups.csv not found. Run matching_test.py first.")

    ttps_df = pd.read_csv("inputted_ttps.csv")
    matched_df = pd.read_csv("matched_groups.csv")

    input_ttps = ttps_df["TTP"].dropna().tolist()
    return input_ttps, matched_df

def analyze_TTP(input_ttps, matched_df):
    mapping_summary = "\n".join([
        f"{row['group_name']}: {row.to_dict()}"
        for _, row in matched_df.head(10).iterrows()
        if "group_name" in row
    ])

    prompt = f"""
You are a cyber threat intelligence analyst.
You are given:
1. A list of matched attacker groups (with details) from a TTP matching engine.
2. A list of detected MITRE ATT&CK TTPs from a security incident.

Your task:
- Compare the detected TTPs with known attacker groups.
- Identify the most likely actor(s) based on overlapping TTPs.
- Explain reasoning clearly and end with a confidence rating.

Matched Actor Groups:
{mapping_summary}

Detected Input TTPs:
{', '.join(input_ttps)}

Return your analysis formatted as:
---
**Analysis Summary:**
[Explanation]

**Overlap Table:**
| Actor Group | Overlap Count |
|--------------|---------------|
| Example APT  | 3 |
| Example B    | 2 |
| Example C    | 1 |

**Most Likely Attacker:**
[Actor Name] – Confidence: [High/Medium/Low]
---
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.3,
        messages=[
            {"role": "system", "content": "You are an experienced MITRE ATT&CK threat analyst."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content.strip()

def parse_ai_response(text):
    sections = {"summary": "", "table": "", "attacker": ""}
    summary_match = re.search(r"\*\*Analysis Summary:\*\*(.*?)\*\*Overlap Table:", text, re.S)
    table_match = re.search(r"\*\*Overlap Table:\*\*(.*?)\*\*Most Likely Attacker:", text, re.S)
    attacker_match = re.search(r"\*\*Most Likely Attacker:\*\*(.*)", text, re.S)

    if summary_match:
        sections["summary"] = summary_match.group(1).strip()
    if table_match:
        sections["table"] = table_match.group(1).strip()
    if attacker_match:
        sections["attacker"] = attacker_match.group(1).strip()

    return sections

def generate_word_report(report_text, input_ttps):
    parsed = parse_ai_response(report_text)
    doc = Document()

    # Title
    title = doc.add_heading("Threat Attribution Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Detected TTPs: {', '.join(input_ttps)}")

    doc.add_paragraph("")

    # Section 1: Analysis Summary
    doc.add_heading("1. Analysis Summary", level=2)
    doc.add_paragraph(parsed["summary"] or "No summary provided.")
    doc.add_paragraph("")

    # Section 2: Overlap Table
    doc.add_heading("2. Overlap Table", level=2)
    if parsed["table"]:
        lines = [line.strip("| ").split("|") for line in parsed["table"].splitlines() if "|" in line]
        lines = [l for l in lines if len(l) == 2 and not l[0].startswith("-")]
        if len(lines) > 1:
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Actor Group"
            hdr_cells[1].text = "Overlap Count"
            for row in lines[1:]:
                row_cells = table.add_row().cells
                row_cells[0].text = row[0].strip()
                row_cells[1].text = row[1].strip()
        else:
            doc.add_paragraph("No overlap data found.")
    else:
        doc.add_paragraph("No overlap table provided.")
    doc.add_paragraph("")

    # Section 3: Most Likely Attacker
    doc.add_heading("3. Most Likely Attacker", level=2)
    doc.add_paragraph(parsed["attacker"] or "No attacker identified.")

    filename = f"Threat_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    print(f"Report saved as {filename}")

if __name__ == "__main__":
    print("Running AI Threat Analysis...")
    input_ttps, matched_df = load_csv_data()
    report = analyze_TTP(input_ttps, matched_df)
    generate_word_report(report, input_ttps)
