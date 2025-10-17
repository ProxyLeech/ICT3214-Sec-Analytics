import os
from typing import Optional
import re
import pandas as pd
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path

# ===========================
# Setup and Data Loading
# ===========================
load_dotenv()
key = os.getenv("OPENAI_API_KEY")
if not key:
    raise SystemExit("OPENAI_API_KEY not found in .env")

client = OpenAI(api_key=key)

def load_csv_data():
    if not os.path.exists("inputted_ttps_rule.csv"):
        raise SystemExit("inputted_ttps_rule.csv not found.")
    if not os.path.exists("matched_groups_rule.csv"):
        raise SystemExit("matched_groups_rule.csv not found.")

    ttps_df = pd.read_csv("inputted_ttps_rule.csv")
    matched_df = pd.read_csv("matched_groups_rule.csv")

    if "score" in matched_df.columns:
        matched_df = matched_df.sort_values(by="score", ascending=False)

    input_ttps = ttps_df["TTP"].dropna().tolist()
    return input_ttps, matched_df


def load_mitigations_summary(mitigations_csv: str) -> str:
    """
    Load and summarize mitigations.csv content to inject directly into the report
    (without GPT generation).
    """
    if not os.path.exists(mitigations_csv):
        return "No mitigations file found."
    
    try:
        df = pd.read_csv(mitigations_csv)
        cols = [c for c in ["target id", "target name", "mapping description"] if c in df.columns]
        if not cols:
            return "Mitigations data is missing expected columns."

        df = df[cols].dropna(how="all")

        lines = []
        for (tid, tname), grp in df.groupby([c for c in ["target id", "target name"] if c in df.columns]):
            desc_col = "mapping description" if "mapping description" in grp.columns else None
            descs = []
            if desc_col:
                descs = grp[desc_col].dropna().astype(str).head(2).tolist()
            line = f"{tid} – {tname}\n" + "\n".join(f"• {d}" for d in descs)
            lines.append(line)
            if len(lines) > 40:
                break

        if not lines:
            return "No mitigation mappings available."

        return "\n".join(lines)

    except Exception as e:
        return f"Error reading mitigations CSV: {e}"
    
def load_filtered_mitigations(mitigations_csv: str, ttps: list[str]) -> pd.DataFrame:
    """
    Return only mitigation rows whose 'target id' matches any TTP in `ttps`,
    including sub-techniques (e.g., T1110.x matches T1110).
    """
    if not os.path.exists(mitigations_csv):
        raise FileNotFoundError(f"{mitigations_csv} not found")

    df = pd.read_csv(mitigations_csv)
    if "target id" not in df.columns:
        raise ValueError("mitigations.csv missing 'target id' column")

    df["target id"] = df["target id"].astype(str).str.strip().str.upper()

    roots = {t.split(".")[0] for t in ttps}
    def _match(tid: str) -> bool:
        tid = tid.upper().strip()
        return any(tid == t or tid.startswith(f"{t}.") for t in ttps) or tid.split(".")[0] in roots

    return df[df["target id"].apply(_match)].copy()


# ===========================
# GPT Analysis
# ===========================
def analyze_TTP(input_ttps, matched_df, mitigations_csv: Optional[str] = None):
    mapping_summary = "\n".join([
        f"{row['group_name']}: {row.to_dict()}"
        for _, row in matched_df.head(10).iterrows()
        if "group_name" in row
    ])

    # NEW: If a mitigations CSV is provided, load & summarize succinctly
    mitigations_block = ""
    if mitigations_csv and os.path.exists(mitigations_csv):
        try:
            mdf = pd.read_csv(mitigations_csv)
            # keep only the canonical columns if present
            cols = [c for c in ["target id", "target name", "mapping description"] if c in mdf.columns]
            if cols:
                mdf = mdf[cols].dropna(how="all")
                # compact: group by technique, take up to 2 example mappings each, limit overall lines
                lines = []
                for (tid, tname), grp in mdf.groupby([c for c in ["target id", "target name"] if c in mdf.columns]):
                    desc_col = "mapping description" if "mapping description" in grp.columns else None
                    examples = []
                    if desc_col:
                        for d in grp[desc_col].dropna().astype(str).head(2).tolist():
                            examples.append(f"- {d[:200]}")
                    header = f"{tid} – {tname}" if isinstance(tid, str) and isinstance(tname, str) else str((tid, tname))
                    block = header if not examples else header + "\n" + "\n".join(examples)
                    lines.append(block)
                    if len(lines) >= 40:  # cap to keep prompt small
                        break
                if lines:
                    mitigations_block = "Associated Mitigations (sample):\n" + "\n".join(lines)
        except Exception as _e:
            # Non-fatal: just skip extra context if anything goes wrong
            mitigations_block = ""

    prompt = f"""
You are a cyber threat intelligence analyst.
You are given:
1. A list of matched attacker groups (with details) from a TTP matching engine.
2. A list of detected MITRE ATT&CK TTPs from a security incident.
3. Additional CSV-derived **associated mitigations** for techniques (if provided).

Your task:
- Compare the detected TTPs with known attacker groups.
- Identify the most likely actor(s) based on overlapping TTPs.
- Weigh the **associated mitigations** context when proposing defensive actions.
- Explain reasoning clearly and end with a confidence rating.

Matched Actor Groups:
{mapping_summary}

Detected Input TTPs:
{', '.join(input_ttps)}

Return your analysis formatted as:
---
**Analysis Summary:**
[Explanation]

**Overlap Table:**
| Actor Group | Overlap Count |
|--------------|---------------|
| Example APT  | 3 |
| Example B    | 2 |
| Example C    | 1 |

**Most Likely Attacker:**
[Actor Name] – Confidence: [High/Medium/Low]
---
"""

    prompt += """
Also include:
- A concise **summary** explaining overall findings and observed behavior patterns.
- A **justification** for why specific attacker groups are likely involved based on technique overlap.
- A section suggesting **defensive mitigations or detections** organizations can apply to counter these TTPs (use the **associated mitigations** context if available).
- Conclude with a short **confidence rating** (High/Medium/Low) and brief reasoning for this rating.

Keep the language concise, analytical, and professional.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.3,
        messages=[
            {"role": "system", "content": "You are an experienced MITRE ATT&CK threat analyst."},
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content.strip()

# ===========================
# Parse Response
# ===========================
# def parse_ai_response(text):
#     """Extract structured sections from GPT response."""
#     sections = {
#         "summary": "", "table": "", "attacker": "",
#         "mitigation": "", "suggestion": ""
#     }

#     summary_match = re.search(r"\*\*Analysis Summary:\*\*(.*?)\*\*Overlap Table:", text, re.S)
#     table_match = re.search(r"\*\*Overlap Table:\*\*(.*?)\*\*Most Likely Attacker:", text, re.S)
#     attacker_match = re.search(r"\*\*Most Likely Attacker:\*\*(.*?)(\*\*Defensive Mitigations|\*\*Mitigations|\*\*Mitigation)", text, re.S)
#     mitigation_match = re.search(r"\*\*Defensive Mitigations.*?\*\*(.*?)\*\*Confidence Rating", text, re.S)
#     suggestion_match = re.search(r"\*\*Confidence Rating.*?\*\*(.*)", text, re.S)

#     if summary_match:
#         sections["summary"] = summary_match.group(1).strip()
#     if table_match:
#         sections["table"] = table_match.group(1).strip()
#     if attacker_match:
#         sections["attacker"] = attacker_match.group(1).strip()
#     if mitigation_match:
#         sections["mitigation"] = mitigation_match.group(1).strip()
#     if suggestion_match:
#         sections["suggestion"] = suggestion_match.group(1).strip()

#     # Strip markdown bold markers (e.g. **text**)
#     for key in sections:
#         sections[key] = re.sub(r"\*\*(.*?)\*\*", r"\1", sections[key])

#     return sections
def parse_ai_response(text: str) -> dict:
    """
    Parse GPT output into sections expected by the HTML templates:
      - summary
      - table
      - attacker
      - mitigation
      - suggestion

    Tolerates multiple heading variants and formats (bold markdown, H3, plain).
    """
    import re

    # Normalize whitespace and keep a working copy
    t = text.replace("\r\n", "\n").strip()

    # Map many possible headings to canonical tokens
    heading_patterns = {
        "SUMMARY": [
            r"\*\*\s*Analysis\s+Summary\s*:\s*\*\*", r"^#{1,6}\s*Analysis\s+Summary\b", r"\bAnalysis\s+Summary\s*:"
        ],
        "TABLE": [
            r"\*\*\s*Overlap\s*Table\s*:\s*\*\*", r"\*\*\s*Technique\s+Overlap\s+Overview\s*:\s*\*\*",
            r"^#{1,6}\s*(Overlap\s*Table|Technique\s+Overlap\s+Overview)\b", r"\b(Overlap\s*Table|Technique\s+Overlap\s+Overview)\s*:"
        ],
        "ATTACKER": [
            r"\*\*\s*Most\s+Likely\s+Attacker\s*:\s*\*\*", r"\*\*\s*Probable\s+Threat\s+Actor\(s\)\s*:\s*\*\*",
            r"^#{1,6}\s*(Most\s+Likely\s+Attacker|Probable\s+Threat\s+Actor\(s\))\b",
            r"\b(Most\s+Likely\s+Attacker|Probable\s+Threat\s+Actor\(s\))\s*:"
        ],
        "MITIGATION": [
            r"\*\*\s*Defensive\s+Mitigations\s*.*?\*\*", r"\*\*\s*Mitigations?\s*.*?\*\*",
            r"^#{1,6}\s*(Defensive\s+Mitigations?|Mitigations?)\b",
            r"\b(Defensive\s+Mitigations?|Mitigations?)\s*:"
        ],
        "SUGGESTION": [
            r"\*\*\s*Analyst\s+Recommendations\s*.*?\*\*", r"\*\*\s*Recommendations\s*.*?\*\*",
            r"\*\*\s*Confidence\s+Rating\s*.*?\*\*",  # sometimes content ends up under this
            r"^#{1,6}\s*(Analyst\s+Recommendations|Recommendations|Confidence\s+Rating)\b",
            r"\b(Analyst\s+Recommendations|Recommendations|Confidence\s+Rating)\s*:"
        ],
    }

    # Build a single alternation that turns any matching heading into a token line
    token_map = {k: f"[[{k}]]" for k in heading_patterns}
    alternations = []
    for pats in heading_patterns.values():
        alternations.extend(pats)
    big_rx = re.compile("(" + "|".join(pats for pats in alternations) + ")", re.IGNORECASE | re.MULTILINE)

    # Substitute matches with canonical tokens
    def _sub(m):
        s = m.group(0)
        for key, pats in heading_patterns.items():
            for p in pats:
                if re.fullmatch(p, s, flags=re.IGNORECASE):
                    return token_map[key]
        # If we got here, it's a partial match – pick first that contains
        for key, pats in heading_patterns.items():
            if any(re.search(p, s, re.IGNORECASE) for p in pats):
                return token_map[key]
        return s

    normalized = big_rx.sub(_sub, t)

    # Split into sections by tokens
    sections = {"summary": "", "table": "", "attacker": "", "mitigation": "", "suggestion": ""}
    order = ["SUMMARY", "TABLE", "ATTACKER", "MITIGATION", "SUGGESTION"]
    # Ensure the first token exists to simplify splitting logic
    if "[[SUMMARY]]" not in normalized:
        normalized = "[[SUMMARY]]\n" + normalized

    # Extract text between tokens
    def grab(block, start_tok, end_tok=None):
        start_idx = block.find(start_tok)
        if start_idx == -1:
            return ""
        start_idx += len(start_tok)
        if end_tok:
            end_idx = block.find(end_tok, start_idx)
            return block[start_idx:end_idx].strip() if end_idx != -1 else block[start_idx:].strip()
        return block[start_idx:].strip()

    cur = normalized
    for i, key in enumerate(order):
        start_tok = f"[[{key}]]"
        end_tok = f"[[{order[i+1]}]]" if i < len(order) - 1 else None
        val = grab(cur, start_tok, end_tok)
        # move window forward
        if end_tok:
            cur = cur[cur.find(end_tok):]
        # save
        k = key.lower()
        if k == "suggestion" and not val:
            # sometimes people dump suggestions under "Confidence Rating"
            pass
        sections_map = {
            "summary": "summary",
            "table": "table",
            "attacker": "attacker",
            "mitigation": "mitigation",
            "suggestion": "suggestion",
        }
        sections[sections_map[k]] = val

    # Strip stray bold markers
    for k, v in sections.items():
        sections[k] = re.sub(r"\*\*(.*?)\*\*", r"\1", v or "").strip()

    return sections

# ===========================
# Generate Word Report
# ===========================
def load_mitigations_summary(mitigations_csv: str) -> str:
    """
    Load and summarize mitigations.csv content for the defensive mitigations section.
    """
    if not os.path.exists(mitigations_csv):
        return "No mitigations file found."

    try:
        df = pd.read_csv(mitigations_csv)
        cols = [c for c in ["target id", "target name", "mapping description"] if c in df.columns]
        if not cols:
            return "Mitigations file missing expected columns."

        df = df[cols].dropna(how="all")

        lines = []
        for (tid, tname), grp in df.groupby([c for c in ["target id", "target name"] if c in df.columns]):
            desc_col = "mapping description" if "mapping description" in grp.columns else None
            descs = []
            if desc_col:
                descs = grp[desc_col].dropna().astype(str).head(2).tolist()
            line = f"{tid} – {tname}\n" + "\n".join(f"• {d}" for d in descs)
            lines.append(line)
            if len(lines) >= 40:
                break

        return "\n".join(lines) if lines else "No mitigation mappings found."

    except Exception as e:
        return f"Error reading mitigations CSV: {e}"

def generate_word_report(report_text, input_ttps, mitigations_csv=None):
    parsed = parse_ai_response(report_text)

    # Resolve base directory relative to this script
    base_dir = Path(__file__).resolve().parent
    template_path = base_dir / "templates" / "cleaned_report_template.docx"

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found at: {template_path}")

    doc = Document(template_path)

    # Update metadata
    for para in doc.paragraphs:
        if "Generated on:" in para.text:
            para.text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        elif "Detected TTPs:" in para.text:
            para.text = f"Detected TTPs: {', '.join(input_ttps)}"

    # Populate main report sections
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        if "1. analysis summary" in text:
            doc.paragraphs[i + 1].text = parsed["summary"] or "N/A"

        elif "2. overlap table" in text:
            lines = [l.strip("| ").split("|") for l in parsed["table"].splitlines() if "|" in l]
            if len(lines) > 1 and len(doc.tables) > 0:
                table = doc.tables[0]
                for row_data in lines[1:]:
                    row = table.add_row().cells
                    row[0].text = row_data[0].strip()
                    row[1].text = row_data[1].strip()

        elif "3. most likely attacker" in text:
            doc.paragraphs[i + 1].text = parsed["attacker"] or "N/A"

        elif "4. defensive mitigations" in text:
            if mitigations_csv and os.path.exists(mitigations_csv):
                mitigations_text = load_mitigations_summary(mitigations_csv)
                doc.paragraphs[i + 1].text = mitigations_text
            else:
                # fallback to parsed mitigation if CSV missing
                doc.paragraphs[i + 1].text = parsed.get("mitigation", "[Mitigations CSV not found or invalid.]")

        elif "5. analyst suggestions" in text:
            doc.paragraphs[i + 1].text = parsed["suggestion"] or "[Add reflection, lessons learned, or recommendations here.]"

    # Save generated file
    ##output_dir = base_dir / "Generated_Reports"
    ##output_dir.mkdir(exist_ok=True)
    ## filepath = output_dir / f"Threat_Report_{datetime.now():%Y%m%d_%H%M%S}.docx"
    # doc.save(filepath)
    # print(f"✅ Report saved: {filepath}")

def summarize_mitigations(mitigations: list[dict]) -> str:
    """
    Cleanly summarize mitigations by technique ID.
    Groups repeated mappings and shortens redundant text.
    """
    if not mitigations:
        return "No mitigations found for these techniques."

    import re
    grouped = {}
    for m in mitigations:
        tid = m.get("target id", "").strip()
        name = m.get("target name", "").strip()
        desc = m.get("mapping description", "").strip()

        # Extract first sentence or concise action
        desc = desc.split(".")[0].strip()
        desc = re.sub(r"\s+", " ", desc)

        if not tid:
            continue
        grouped.setdefault(tid, {"name": name, "actions": set()})
        if desc:
            grouped[tid]["actions"].add(desc)

    summarized = []
    for tid, info in grouped.items():
        actions = sorted(info["actions"])
        block = f"{tid} – {info['name']}\n" + "\n".join(f"• {a}" for a in actions)
        summarized.append(block)

    return "\n\n".join(summarized)


# ===========================
# Run
# ===========================
if __name__ == "__main__":
    print("Running AI Threat Analysis...")
    input_ttps, matched_df = load_csv_data()
    report = analyze_TTP(input_ttps, matched_df)
    generate_word_report(report, input_ttps)
