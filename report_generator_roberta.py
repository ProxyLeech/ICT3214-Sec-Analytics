import os
from typing import Optional
import re
import pandas as pd
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from datetime import datetime
from pathlib import Path

# ===========================
# Setup and Data Loading
# ===========================
load_dotenv()
key = os.getenv("OPENAI_API_KEY")
if not key:
    raise SystemExit("OPENAI_API_KEY not found in .env")

client = OpenAI(api_key=key)

def load_csv_data():
    # --- RoBERTa filenames ---
    ttps_csv    = "inputted_ttps.csv"
    matched_csv = "matched_groups_roberta.csv"

    if not os.path.exists(ttps_csv):
        raise SystemExit(f"{ttps_csv} not found.")
    if not os.path.exists(matched_csv):
        raise SystemExit(f"{matched_csv} not found.")

    ttps_df = pd.read_csv(ttps_csv)
    matched_df = pd.read_csv(matched_csv)

    # Normalize columns just in case
    if "group" in matched_df.columns and "group_name" not in matched_df.columns:
        matched_df.rename(columns={"group": "group_name"}, inplace=True)

    if "score" in matched_df.columns:
        matched_df = matched_df.sort_values(by="score", ascending=False)

    input_ttps = ttps_df["TTP"].dropna().tolist()
    return input_ttps, matched_df


def load_mitigations_summary(mitigations_csv: str) -> str:
    if not os.path.exists(mitigations_csv):
        return "No mitigations file found."
    try:
        df = pd.read_csv(mitigations_csv)
        cols = [c for c in ["target id", "target name", "mapping description"] if c in df.columns]
        if not cols:
            return "Mitigations data is missing expected columns."
        df = df[cols].dropna(how="all")

        lines = []
        for (tid, tname), grp in df.groupby([c for c in ["target id", "target name"] if c in df.columns]):
            desc_col = "mapping description" if "mapping description" in grp.columns else None
            descs = grp[desc_col].dropna().astype(str).head(2).tolist() if desc_col else []
            line = f"{tid} – {tname}\n" + "\n".join(f"• {d}" for d in descs)
            lines.append(line)
            if len(lines) >= 40:
                break
        return "\n".join(lines) if lines else "No mitigation mappings available."
    except Exception as e:
        return f"Error reading mitigations CSV: {e}"

# ===========================
# GPT Analysis
# ===========================
def analyze_TTP(input_ttps, matched_df, mitigations_csv: Optional[str] = None):
    mapping_summary = "\n".join([
        f"{row['group_name']}: {row.to_dict()}"
        for _, row in matched_df.head(10).iterrows()
        if "group_name" in row
    ])

    mitigations_block = ""
    if mitigations_csv and os.path.exists(mitigations_csv):
        try:
            mdf = pd.read_csv(mitigations_csv)
            cols = [c for c in ["target id", "target name", "mapping description"] if c in mdf.columns]
            if cols:
                mdf = mdf[cols].dropna(how="all")
                lines = []
                for (tid, tname), grp in mdf.groupby([c for c in ["target id", "target name"] if c in mdf.columns]):
                    desc_col = "mapping description" if "mapping description" in grp.columns else None
                    examples = []
                    if desc_col:
                        for d in grp[desc_col].dropna().astype(str).head(2).tolist():
                            examples.append(f"- {d[:200]}")
                    header = f"{tid} – {tname}" if isinstance(tid, str) and isinstance(tname, str) else str((tid, tname))
                    block = header if not examples else header + "\n" + "\n".join(examples)
                    lines.append(block)
                    if len(lines) >= 40:
                        break
                if lines:
                    mitigations_block = "Associated Mitigations (sample):\n" + "\n".join(lines)
        except Exception:
            mitigations_block = ""

    prompt = f"""
You are a cyber threat intelligence analyst.
You are given:
1. A list of matched attacker groups (with details) from a TTP matching engine (RoBERTa output).
2. A list of detected MITRE ATT&CK TTPs from a security incident.
3. Additional CSV-derived associated mitigations for techniques (if provided).

Your task:
- Compare the detected TTPs with known attacker groups.
- Identify the most likely actor(s) based on overlapping TTPs.
- Weigh the associated mitigations context when proposing defensive actions.
- Explain reasoning clearly and end with a confidence rating.

Matched Actor Groups:
{mapping_summary}

Detected Input TTPs:
{', '.join(input_ttps)}

{mitigations_block}
Return your analysis formatted as:
---
**Analysis Summary:**
[Explanation]

**Overlap Table:**
| Actor Group | Overlap Count |
|--------------|---------------|
| Example APT  | 3 |
| Example B    | 2 |
| Example C    | 1 |

**Most Likely Attacker:**
[Actor Name] – Confidence: [High/Medium/Low]
---
"""

    prompt += """
Also include:
- A concise summary of overall findings and observed behavior patterns.
- A justification for why specific attacker groups are likely involved based on technique overlap.
- A section suggesting defensive mitigations or detections (use the associated mitigations if available).
- Conclude with a short confidence rating (High/Medium/Low) and the reason.

Keep the language concise, analytical, and professional.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.3,
        messages=[
            {"role": "system", "content": "You are an experienced MITRE ATT&CK threat analyst."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content.strip()

# ===========================
# Parse Response (unchanged)
# ===========================
def parse_ai_response(text: str) -> dict:
    import re
    t = text.replace("\r\n", "\n").strip()
    heading_patterns = {
        "SUMMARY": [r"\*\*\s*Analysis\s+Summary\s*:\s*\*\*", r"^#{1,6}\s*Analysis\s+Summary\b", r"\bAnalysis\s+Summary\s*:"],
        "TABLE":   [r"\*\*\s*Overlap\s*Table\s*:\s*\*\*", r"\*\*\s*Technique\s+Overlap\s+Overview\s*:\s*\*", r"^#{1,6}\s*(Overlap\s*Table|Technique\s+Overlap\s+Overview)\b", r"\b(Overlap\s*Table|Technique\s+Overlap\s+Overview)\s*:"],
        "ATTACKER":[r"\*\*\s*Most\s+Likely\s+Attacker\s*:\s*\*\*", r"\*\*\s*Probable\s+Threat\s+Actor\(s\)\s*:\s*\*\*", r"^#{1,6}\s*(Most\s+Likely\s+Attacker|Probable\s+Threat\s+Actor\(s\))\b", r"\b(Most\s+Likely\s+Attacker|Probable\s+Threat\s+Actor\(s\))\s*:"],
        "MITIGATION":[r"\*\*\s*Defensive\s+Mitigations\s*.*?\*\*", r"\*\*\s*Mitigations?\s*.*?\*\*", r"^#{1,6}\s*(Defensive\s+Mitigations?|Mitigations?)\b", r"\b(Defensive\s+Mitigations?|Mitigations?)\s*:"],
        "SUGGESTION":[r"\*\*\s*Analyst\s+Recommendations\s*.*?\*\*", r"\*\*\s*Recommendations\s*.*?\*\*", r"\*\*\s*Confidence\s+Rating\s*.*?\*\*", r"^#{1,6}\s*(Analyst\s+Recommendations|Recommendations|Confidence\s+Rating)\b", r"\b(Analyst\s+Recommendations|Recommendations|Confidence\s+Rating)\s*:"],
    }
    token_map = {k: f"[[{k}]]" for k in heading_patterns}
    alternations = [p for ps in heading_patterns.values() for p in ps]
    big_rx = re.compile("(" + "|".join(alternations) + ")", re.IGNORECASE | re.MULTILINE)

    def _sub(m):
        s = m.group(0)
        for key, pats in heading_patterns.items():
            for p in pats:
                if re.fullmatch(p, s, flags=re.IGNORECASE):
                    return token_map[key]
        for key, pats in heading_patterns.items():
            if any(re.search(p, s, re.IGNORECASE) for p in pats):
                return token_map[key]
        return s

    normalized = big_rx.sub(_sub, t)
    sections = {"summary": "", "table": "", "attacker": "", "mitigation": "", "suggestion": ""}
    order = ["SUMMARY", "TABLE", "ATTACKER", "MITIGATION", "SUGGESTION"]
    if "[[SUMMARY]]" not in normalized:
        normalized = "[[SUMMARY]]\n" + normalized

    def grab(block, start_tok, end_tok=None):
        start_idx = block.find(start_tok)
        if start_idx == -1: return ""
        start_idx += len(start_tok)
        if end_tok:
            end_idx = block.find(end_tok, start_idx)
            return block[start_idx:end_idx].strip() if end_idx != -1 else block[start_idx:].strip()
        return block[start_idx:].strip()

    cur = normalized
    for i, key in enumerate(order):
        start_tok = f"[[{key}]]"
        end_tok = f"[[{order[i+1]}]]" if i < len(order) - 1 else None
        val = grab(cur, start_tok, end_tok)
        if end_tok:
            cur = cur[cur.find(end_tok):]
        sections[key.lower()] = re.sub(r"\*\*(.*?)\*\*", r"\1", val or "").strip()
    return sections

# ===========================
# Generate Word Report
# ===========================
def generate_word_report(report_text, input_ttps, mitigations_csv=None) -> Path:
    parsed = parse_ai_response(report_text)
    base_dir = Path(__file__).resolve().parent
    template_path = base_dir / "templates" / "cleaned_report_template.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found at: {template_path}")

    doc = Document(template_path)

    # Update metadata
    for para in doc.paragraphs:
        if "Generated on:" in para.text:
            para.text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        elif "Detected TTPs:" in para.text:
            para.text = f"Detected TTPs: {', '.join(input_ttps)}"

    # Populate sections
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        if "1. analysis summary" in text:
            doc.paragraphs[i + 1].text = parsed.get("summary", "N/A")

        elif "2. overlap table" in text:
            lines = [l.strip("| ").split("|") for l in parsed.get("table","").splitlines() if "|" in l]
            if len(lines) > 1 and len(doc.tables) > 0:
                table = doc.tables[0]
                for row_data in lines[1:]:
                    row = table.add_row().cells
                    row[0].text = row_data[0].strip()
                    row[1].text = row_data[1].strip()

        elif "3. most likely attacker" in text:
            doc.paragraphs[i + 1].text = parsed.get("attacker", "N/A")

        elif "4. defensive mitigations" in text:
            if mitigations_csv and os.path.exists(mitigations_csv):
                doc.paragraphs[i + 1].text = load_mitigations_summary(mitigations_csv)
            else:
                doc.paragraphs[i + 1].text = parsed.get("mitigation", "[Mitigations CSV not found or invalid.]")

        elif "5. analyst suggestions" in text:
            doc.paragraphs[i + 1].text = parsed.get("suggestion", "[Add reflection, lessons learned, or recommendations here.]")

    # Save generated file
    output_dir = base_dir / "Generated_Reports"
    output_dir.mkdir(exist_ok=True)
    filepath = output_dir / f"Threat_Report_RoBERTa_{datetime.now():%Y%m%d_%H%M%S}.docx"
    doc.save(filepath)
    return filepath

# ===========================
# Run (standalone)
# ===========================
if __name__ == "__main__":
    print("Running AI Threat Analysis (RoBERTa)…")
    input_ttps, matched_df = load_csv_data()
    report = analyze_TTP(input_ttps, matched_df)
    path = generate_word_report(report, input_ttps)
    print(f"✅ Report saved: {path}")
