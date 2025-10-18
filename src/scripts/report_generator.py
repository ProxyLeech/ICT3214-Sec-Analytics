import os
from typing import Optional
import re
import pandas as pd
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
import sys
ROOT = Path(__file__).resolve().parents[2]  # repo root
sys.path.insert(0, str(ROOT))
from project_paths import (
    PROJECT_ROOT, DATA_ROOT, EXPERIMENTS_ROOT, SRC_ROOT, MODELS_ROOT, EXPERIMENTS_ROOT,SCRIPTS_DIR,
    RAW_DIR, PROCESSED_DIR, EXTRACTED_PDFS_DIR,
    MAPPED_DIR, EXCEL_DIR, MITIGATIONS_DIR,
    ATTACK_STIX_DIR,PDFS_DIR,RULES_DIR,EXTRACT_SCRIPT,ATTACK_SCRIPT,MAP_IOCS_SCRIPT,
    BUILD_DATASET_SCRIPT,MITIGATIONS_SCRIPT,
    GROUP_TTPS_DETAIL_CSV,MATCHING_SCRIPT,REPORT_GENERATION_SCRIPT,TECHNIQUE_LABELS_SCRIPT,
    TRAIN_ROBERTA_SCRIPT,PREDICT_SCRIPT,BEST_MODEL_DIR,
    MAPPING_CSV,MITIGATIONS_CSV,EXCEL_ATTACK_TECHS,
    EXTRACTED_IOCS_CSV,TI_GROUPS_TECHS_CSV,DATASET_CSV,LABELS_TXT,GROUP_TTPS_DETAIL_CSV,RANKED_GROUPS_CSV,
    output_dir_for_folds, project_path,ensure_dir_tree,add_src_to_syspath
)
INPUT_TTPS_CSV       = PROJECT_ROOT / "inputted_ttps.csv"
MATCHED_RULE_CSV     = PROJECT_ROOT / "matched_groups_rule.csv"
MATCHED_ROBERTA_CSV  = PROJECT_ROOT / "matched_groups_roberta.csv"
MITIGATIONS_RULE_CSV     = PROJECT_ROOT / "mitigations_rule.csv"
MITIGATIONS_ROBERTA_CSV  = PROJECT_ROOT / "mitigations_roberta.csv"

# ===========================
# Setup and Data Loading
# ===========================
load_dotenv()
key = os.getenv("OPENAI_API_KEY")
if not key:
    raise SystemExit("OPENAI_API_KEY not found in .env")

client = OpenAI(api_key=key)

def load_csv_data():
    for p in (INPUT_TTPS_CSV, MATCHED_RULE_CSV, MATCHED_ROBERTA_CSV):
        if not p.exists():
            raise SystemExit(f"{p} not found.")

    ttps_df            = pd.read_csv(INPUT_TTPS_CSV)
    matched_df_rule    = pd.read_csv(MATCHED_RULE_CSV)
    matched_df_roberta = pd.read_csv(MATCHED_ROBERTA_CSV)

    if "score" in matched_df_rule.columns:
        matched_df_rule = matched_df_rule.sort_values(by="score", ascending=False)
    if "score" in matched_df_roberta.columns:
        matched_df_roberta = matched_df_roberta.sort_values(by="score", ascending=False)

    input_ttps = ttps_df["TTP"].dropna().tolist()
    return input_ttps, matched_df_rule, matched_df_roberta



def load_mitigations_summary(mitigations_csv: str) -> str:
    """
    Load and summarize mitigations.csv content to inject directly into the report
    (without GPT generation).
    """
    if not os.path.exists(mitigations_csv):
        return "No mitigations file found."
    
    try:
        df = MITIGATIONS_CSV
        cols = [c for c in ["target id", "target name", "mapping description"] if c in df.columns]
        if not cols:
            return "Mitigations data is missing expected columns."

        df = df[cols].dropna(how="all")

        lines = []
        for (tid, tname), grp in df.groupby([c for c in ["target id", "target name"] if c in df.columns]):
            desc_col = "mapping description" if "mapping description" in grp.columns else None
            descs = []
            if desc_col:
                descs = grp[desc_col].dropna().astype(str).head(2).tolist()
            line = f"{tid} – {tname}\n" + "\n".join(f"• {d}" for d in descs)
            lines.append(line)
            if len(lines) > 40:
                break

        if not lines:
            return "No mitigation mappings available."

        return "\n".join(lines)

    except Exception as e:
        return f"Error reading mitigations CSV: {e}"
    
def load_filtered_mitigations(mitigations_csv: str, ttps: list[str]) -> pd.DataFrame:
    """
    Return only mitigation rows whose 'target id' matches any TTP in `ttps`,
    including sub-techniques (e.g., T1110.x matches T1110).
    """
    if not os.path.exists(mitigations_csv):
        raise FileNotFoundError(f"{mitigations_csv} not found")

    df = pd.read_csv(mitigations_csv)
    if "target id" not in df.columns:
        raise ValueError("mitigations.csv missing 'target id' column")

    df["target id"] = df["target id"].astype(str).str.strip().str.upper()

    roots = {t.split(".")[0] for t in ttps}
    def _match(tid: str) -> bool:
        tid = tid.upper().strip()
        return any(tid == t or tid.startswith(f"{t}.") for t in ttps) or tid.split(".")[0] in roots

    return df[df["target id"].apply(_match)].copy()


# ===========================
# GPT Analysis
# ===========================
def analyze_TTP(input_ttps, matched_df, mitigations_csv: Optional[str] = None):
    mapping_summary = "\n".join([
        f"{row['group_name']}: {row.to_dict()}"
        for _, row in matched_df.head(10).iterrows()
        if "group_name" in row
    ])

    # NEW: If a mitigations CSV is provided, load & summarize succinctly
    mitigations_block = ""
    if mitigations_csv and os.path.exists(mitigations_csv):
        try:
            mdf = pd.read_csv(mitigations_csv)
            # keep only the canonical columns if present
            cols = [c for c in ["target id", "target name", "mapping description"] if c in mdf.columns]
            if cols:
                mdf = mdf[cols].dropna(how="all")
                # compact: group by technique, take up to 2 example mappings each, limit overall lines
                lines = []
                for (tid, tname), grp in mdf.groupby([c for c in ["target id", "target name"] if c in mdf.columns]):
                    desc_col = "mapping description" if "mapping description" in grp.columns else None
                    examples = []
                    if desc_col:
                        for d in grp[desc_col].dropna().astype(str).head(2).tolist():
                            examples.append(f"- {d[:200]}")
                    header = f"{tid} – {tname}" if isinstance(tid, str) and isinstance(tname, str) else str((tid, tname))
                    block = header if not examples else header + "\n" + "\n".join(examples)
                    lines.append(block)
                    if len(lines) >= 40:  # cap to keep prompt small
                        break
                if lines:
                    mitigations_block = "Associated Mitigations (sample):\n" + "\n".join(lines)
        except Exception as _e:
            # Non-fatal: just skip extra context if anything goes wrong
            mitigations_block = ""

    prompt = f"""
You are a cyber threat intelligence analyst.
You are given:
1. A list of matched attacker groups (with details) from a TTP matching engine.
2. A list of detected MITRE ATT&CK TTPs from a security incident.
3. Additional CSV-derived **associated mitigations** for techniques (if provided).

Your task:
- Compare the detected TTPs with known attacker groups.
- Identify the most likely actor(s) based on overlapping TTPs.
- Weigh the **associated mitigations** context when proposing defensive actions.
- Explain reasoning clearly and end with a confidence rating.

Matched Actor Groups:
{mapping_summary}

Detected Input TTPs:
{', '.join(input_ttps)}

Return your analysis formatted as:
---
**Analysis Summary:**
[Explanation]

**Overlap Table:**
| Actor Group | Overlap Count |
|--------------|---------------|
| Example APT  | 3 |
| Example B    | 2 |
| Example C    | 1 |

**Most Likely Attacker:**
[Actor Name] – Confidence: [High/Medium/Low]
---
"""

    prompt += """
Also include:
- A concise **summary** explaining overall findings and observed behavior patterns.
- A **justification** for why specific attacker groups are likely involved based on technique overlap.
- A section suggesting **defensive mitigations or detections** organizations can apply to counter these TTPs (use the **associated mitigations** context if available).
- Conclude with a short **confidence rating** (High/Medium/Low) and brief reasoning for this rating.

Keep the language concise, analytical, and professional.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.3,
        messages=[
            {"role": "system", "content": "You are an experienced MITRE ATT&CK threat analyst."},
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content.strip()

# ===========================
# Parse Response
# ===========================
def parse_ai_response(text: str) -> dict:
    """
    Parse GPT output into sections expected by the HTML templates:
      - summary
      - table
      - attacker
      - mitigation
      - suggestion

    Tolerates multiple heading variants and formats (bold markdown, H3, plain).
    """
    import re

    # Normalize whitespace and keep a working copy
    t = text.replace("\r\n", "\n").strip()

    # Map many possible headings to canonical tokens
    heading_patterns = {
        "SUMMARY": [
            r"\*\*\s*Analysis\s+Summary\s*:\s*\*\*", r"^#{1,6}\s*Analysis\s+Summary\b", r"\bAnalysis\s+Summary\s*:"
        ],
        "TABLE": [
            r"\*\*\s*Overlap\s*Table\s*:\s*\*\*", r"\*\*\s*Technique\s+Overlap\s+Overview\s*:\s*\*\*",
            r"^#{1,6}\s*(Overlap\s*Table|Technique\s+Overlap\s+Overview)\b", r"\b(Overlap\s*Table|Technique\s+Overlap\s+Overview)\s*:"
        ],
        "ATTACKER": [
            r"\*\*\s*Most\s+Likely\s+Attacker\s*:\s*\*\*", r"\*\*\s*Probable\s+Threat\s+Actor\(s\)\s*:\s*\*\*",
            r"^#{1,6}\s*(Most\s+Likely\s+Attacker|Probable\s+Threat\s+Actor\(s\))\b",
            r"\b(Most\s+Likely\s+Attacker|Probable\s+Threat\s+Actor\(s\))\s*:"
        ],
        "MITIGATION": [
            r"\*\*\s*Defensive\s+Mitigations\s*.*?\*\*", r"\*\*\s*Mitigations?\s*.*?\*\*",
            r"^#{1,6}\s*(Defensive\s+Mitigations?|Mitigations?)\b",
            r"\b(Defensive\s+Mitigations?|Mitigations?)\s*:"
        ],
        "SUGGESTION": [
            r"\*\*\s*Analyst\s+Recommendations\s*.*?\*\*", r"\*\*\s*Recommendations\s*.*?\*\*",
            r"\*\*\s*Confidence\s+Rating\s*.*?\*\*",  # sometimes content ends up under this
            r"^#{1,6}\s*(Analyst\s+Recommendations|Recommendations|Confidence\s+Rating)\b",
            r"\b(Analyst\s+Recommendations|Recommendations|Confidence\s+Rating)\s*:"
        ],
    }

    # Build a single alternation that turns any matching heading into a token line
    token_map = {k: f"[[{k}]]" for k in heading_patterns}
    alternations = []
    for pats in heading_patterns.values():
        alternations.extend(pats)
    big_rx = re.compile("(" + "|".join(pats for pats in alternations) + ")", re.IGNORECASE | re.MULTILINE)

    # Substitute matches with canonical tokens
    def _sub(m):
        s = m.group(0)
        for key, pats in heading_patterns.items():
            for p in pats:
                if re.fullmatch(p, s, flags=re.IGNORECASE):
                    return token_map[key]
        # If we got here, it's a partial match – pick first that contains
        for key, pats in heading_patterns.items():
            if any(re.search(p, s, re.IGNORECASE) for p in pats):
                return token_map[key]
        return s

    normalized = big_rx.sub(_sub, t)

    # Split into sections by tokens
    sections = {"summary": "", "table": "", "attacker": "", "mitigation": "", "suggestion": ""}
    order = ["SUMMARY", "TABLE", "ATTACKER", "MITIGATION", "SUGGESTION"]
    # Ensure the first token exists to simplify splitting logic
    if "[[SUMMARY]]" not in normalized:
        normalized = "[[SUMMARY]]\n" + normalized

    # Extract text between tokens
    def grab(block, start_tok, end_tok=None):
        start_idx = block.find(start_tok)
        if start_idx == -1:
            return ""
        start_idx += len(start_tok)
        if end_tok:
            end_idx = block.find(end_tok, start_idx)
            return block[start_idx:end_idx].strip() if end_idx != -1 else block[start_idx:].strip()
        return block[start_idx:].strip()

    cur = normalized
    for i, key in enumerate(order):
        start_tok = f"[[{key}]]"
        end_tok = f"[[{order[i+1]}]]" if i < len(order) - 1 else None
        val = grab(cur, start_tok, end_tok)
        # move window forward
        if end_tok:
            cur = cur[cur.find(end_tok):]
        # save
        k = key.lower()
        if k == "suggestion" and not val:
            # sometimes people dump suggestions under "Confidence Rating"
            pass
        sections_map = {
            "summary": "summary",
            "table": "table",
            "attacker": "attacker",
            "mitigation": "mitigation",
            "suggestion": "suggestion",
        }
        sections[sections_map[k]] = val

    # Strip stray bold markers
    for k, v in sections.items():
        sections[k] = re.sub(r"\*\*(.*?)\*\*", r"\1", v or "").strip()

    return sections

# ===========================
# Generate Word Report
# ===========================
def load_mitigations_summary(mitigations_csv: str) -> str:
    """
    Load and summarize mitigations.csv content for the defensive mitigations section.
    """
    if not os.path.exists(mitigations_csv):
        return "No mitigations file found."

    try:
        df = pd.read_csv(mitigations_csv)
        cols = [c for c in ["target id", "target name", "mapping description"] if c in df.columns]
        if not cols:
            return "Mitigations file missing expected columns."

        df = df[cols].dropna(how="all")

        lines = []
        for (tid, tname), grp in df.groupby([c for c in ["target id", "target name"] if c in df.columns]):
            desc_col = "mapping description" if "mapping description" in grp.columns else None
            descs = []
            if desc_col:
                descs = grp[desc_col].dropna().astype(str).head(2).tolist()
            line = f"{tid} – {tname}\n" + "\n".join(f"• {d}" for d in descs)
            lines.append(line)
            if len(lines) >= 40:
                break

        return "\n".join(lines) if lines else "No mitigation mappings found."

    except Exception as e:
        return f"Error reading mitigations CSV: {e}"

def generate_word_report(report_text, input_ttps, mitigations_csv=None):
    parsed = parse_ai_response(report_text)

    template_path = PROJECT_ROOT / "templates" / "cleaned_report_template.docx"

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found at: {template_path}")

    doc = Document(template_path)

    # Update metadata
    for para in doc.paragraphs:
        if "Generated on:" in para.text:
            para.text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        elif "Detected TTPs:" in para.text:
            para.text = f"Detected TTPs: {', '.join(input_ttps)}"

    # Populate main report sections
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        if "1. analysis summary" in text:
            doc.paragraphs[i + 1].text = parsed["summary"] or "N/A"

        elif "2. overlap table" in text:
            lines = [l.strip("| ").split("|") for l in parsed["table"].splitlines() if "|" in l]
            if len(lines) > 1 and len(doc.tables) > 0:
                table = doc.tables[0]
                for row_data in lines[1:]:
                    row = table.add_row().cells
                    row[0].text = row_data[0].strip()
                    row[1].text = row_data[1].strip()

        elif "3. most likely attacker" in text:
            doc.paragraphs[i + 1].text = parsed["attacker"] or "N/A"

        elif "4. defensive mitigations" in text:
            if mitigations_csv and os.path.exists(mitigations_csv := str(mitigations_csv)):
                mitigations_text = load_mitigations_summary(mitigations_csv)
                doc.paragraphs[i + 1].text = mitigations_text
            else:
                # fallback to parsed mitigation if CSV missing
                doc.paragraphs[i + 1].text = parsed.get("mitigation", "[Mitigations CSV not found or invalid.]")

        elif "5. analyst suggestions" in text:
            doc.paragraphs[i + 1].text = parsed["suggestion"] or "[Add reflection, lessons learned, or recommendations here.]"

def summarize_mitigations(mitigations: list[dict]) -> str:
    """
    Cleanly summarize mitigations by technique ID.
    Groups repeated mappings and shortens redundant text.
    """
    if not mitigations:
        return "No mitigations found for these techniques."

    import re
    grouped = {}
    for m in mitigations:
        tid = m.get("target id", "").strip()
        name = m.get("target name", "").strip()
        desc = m.get("mapping description", "").strip()

        # Extract first sentence or concise action
        desc = desc.split(".")[0].strip()
        desc = re.sub(r"\s+", " ", desc)

        if not tid:
            continue
        grouped.setdefault(tid, {"name": name, "actions": set()})
        if desc:
            grouped[tid]["actions"].add(desc)

    summarized = []
    for tid, info in grouped.items():
        actions = sorted(info["actions"])
        block = f"{tid} – {info['name']}\n" + "\n".join(f"• {a}" for a in actions)
        summarized.append(block)

    return "\n\n".join(summarized)
def _filter_and_write_mitigations(mitigations_csv: str, ttps: list[str], out_path: str) -> Optional[str]:
    """
    Write a mitigations CSV filtered to the provided TTP list.
    Returns the CSV path if something was written; otherwise None.
    """
    try:
        df = load_filtered_mitigations(mitigations_csv, ttps)
        if not df.empty:
            pd_path = Path(out_path)
            pd_path.parent.mkdir(parents=True, exist_ok=True)
            df.to_csv(pd_path, index=False)
            return str(pd_path)
    except Exception as _e:
        pass
    return None

def _extract_ttps_from_text(s: str) -> list[str]:
    """Find T####(.###) technique IDs in a string."""
    if not isinstance(s, str):
        return []
    return re.findall(r"\bT\d{4}(?:\.\d{3})?\b", s, flags=re.IGNORECASE)

def _load_top_group_ttps(matched_df: pd.DataFrame) -> list[str]:
    """
    Return a sorted unique list of technique IDs mapped to the FIRST (top-scored) group
    in matched_df using Data/mapped/group_ttps_detail.csv. Falls back to [].
    """
    # Pick the top row/group robustly
    if matched_df is None or matched_df.empty:
        return []
    # Prefer highest score if present
    df_sorted = matched_df.copy()
    if "score" in df_sorted.columns:
        df_sorted = df_sorted.sort_values("score", ascending=False)
    top = df_sorted.iloc[0]

    top_group_name = None
    top_group_id = None
    for c in ("group_name", "group", "actor", "name"):
        if c in df_sorted.columns:
            top_group_name = str(top[c]).strip()
            break
    for c in ("group_id", "id", "mitre_id"):
        if c in df_sorted.columns:
            top_group_id = str(top[c]).strip()
            break


    map_csv = GROUP_TTPS_DETAIL_CSV
    if not map_csv.exists():
        # No strict map -> no top-group-filtered ttps
        return []

    try:
        g = pd.read_csv(map_csv)
    except Exception:
        return []

    # Normalize columns
    g.columns = [c.strip().lower() for c in g.columns]
    name_col = "group_name" if "group_name" in g.columns else None
    id_col   = "group_id"   if "group_id"   in g.columns else None

    # Filter rows belonging to the top group by id or name
    sel = pd.Series([True] * len(g))
    if top_group_id and id_col:
        sel &= g[id_col].astype(str).str.strip().str.lower() == top_group_id.strip().lower()
    elif top_group_name and name_col:
        sel &= g[name_col].astype(str).str.strip().str.lower() == top_group_name.strip().lower()
    else:
        return []

    gsel = g[sel]
    if gsel.empty:
        return []

    # Collect TTPs from known text columns
    ttps = []
    for col in ("matched_exact", "matched_root_only", "ttp_list", "techniques", "technique_ids"):
        if col in gsel.columns:
            gsel[col] = gsel[col].fillna("")
            for txt in gsel[col].tolist():
                ttps.extend(_extract_ttps_from_text(txt))

    # De-duplicate and normalize to uppercase; sort by numeric order
    ttps = {t.upper() for t in ttps}
    def _key(tid: str):
        m = re.match(r"T(\d{4})(?:\.(\d{3}))?$", tid, re.IGNORECASE)
        return (int(m.group(1)), int(m.group(2) or 999)) if m else (9999, 999)
    return sorted(ttps, key=_key)

def _filter_and_write_mitigations(mitigations_csv: str, ttps: list[str], out_path: str) -> Optional[str]:
    """
    Write a mitigations CSV filtered to the provided TTP list.
    Returns the CSV path if something was written; otherwise None.
    """
    try:
        df = load_filtered_mitigations(mitigations_csv, ttps)
        if not df.empty:
            pd_path = Path(out_path)
            pd_path.parent.mkdir(parents=True, exist_ok=True)
            df.to_csv(pd_path, index=False)
            return str(pd_path)
    except Exception:
        pass
    return None


# ===========================
# Run
# ===========================
if __name__ == "__main__":
    print("Running AI Threat Analysis for RULES and ROBERTA (top group mitigations only)…")

    # --- Input CSVs you specified ---
    try:
        ttps_df            = pd.read_csv(INPUT_TTPS_CSV)
    except FileNotFoundError:
        raise SystemExit("inputted_ttps.csv not found.")

    try:
        matched_df_rules   = pd.read_csv(MATCHED_RULE_CSV)
    except FileNotFoundError:
        raise SystemExit("matched_groups_rule.csv not found.")

    try:
        matched_df_roberta = pd.read_csv(MATCHED_ROBERTA_CSV)
    except FileNotFoundError:
        raise SystemExit("matched_groups_roberta.csv not found.")

    # Extract the selected TTPs (still used in the report header etc.)
    input_ttps = ttps_df["TTP"].dropna().astype(str).str.upper().tolist()

    # Sort matches for a clean prompt (if score is available)
    if "score" in matched_df_rules.columns:
        matched_df_rules = matched_df_rules.sort_values(by="score", ascending=False)
    if "score" in matched_df_roberta.columns:
        matched_df_roberta = matched_df_roberta.sort_values(by="score", ascending=False)

    # --- Mitigations source (global) ---
    mitigations_src = MITIGATIONS_CSV
    if not mitigations_src.exists():
        print("[WARN] Data/mitigations/mitigations.csv not found — proceeding without mitigations context.")
        mit_rules_csv = None
        mit_roberta_csv = None
    else:
        # NEW: limit mitigations to TTPs of the TOP GROUP ONLY for each run
        ttps_rules_top   = _load_top_group_ttps(matched_df_rules)
        ttps_roberta_top = _load_top_group_ttps(matched_df_roberta)

        if not ttps_rules_top:
            print("[WARN] Could not resolve top-group TTPs for RULES; mitigations will be omitted.")
        if not ttps_roberta_top:
            print("[WARN] Could not resolve top-group TTPs for ROBERTA; mitigations will be omitted.")
        mit_rules_csv   = _filter_and_write_mitigations(str(mitigations_src), ttps_rules_top,   str(MITIGATIONS_RULE_CSV))
        mit_roberta_csv = _filter_and_write_mitigations(str(mitigations_src), ttps_roberta_top, str(MITIGATIONS_ROBERTA_CSV))


    # --- RULES analysis+report (mitigations limited to its top group) ---
    print("\n[RUN] RULE-BASED analysis…")
    report_rules = analyze_TTP(input_ttps, matched_df_rules, mitigations_csv=mit_rules_csv)
    generate_word_report(report_rules, input_ttps, mitigations_csv=mit_rules_csv)

    # --- ROBERTA analysis+report (mitigations limited to its top group) ---
    print("\n[RUN] ROBERTA analysis…")
    report_roberta = analyze_TTP(input_ttps, matched_df_roberta, mitigations_csv=mit_roberta_csv)
    generate_word_report(report_roberta, input_ttps, mitigations_csv=mit_roberta_csv)

    print("\n[OK] Done. Generated reports with mitigations scoped to the first actor only (per run).")
