import os
import json
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from datetime import datetime

# Load environment variables
load_dotenv()
key = os.getenv("OPENAI_API_KEY")

if not key:
    raise SystemExit("❌ OPENAI_API_KEY not found in .env!")

client = OpenAI(api_key=key)

def analyze_TTP(mapped_ttp_path, input_ttps):
    """Compare detected TTPs to known attacker mappings using OpenAI reasoning."""
    
    # Load known attacker mappings from JSON
    with open(mapped_ttp_path, "r", encoding="utf-8") as f:
        mapped_data = json.load(f)

    # Summarize mapping for the prompt
    summarized_mapping = "\n".join([
        f"{actor}: {', '.join(ttps)}" for actor, ttps in mapped_data.items()
    ])

    # Build the prompt
    prompt = f"""
You are a cyber threat intelligence analyst.
You will receive:
1. Known MITRE ATT&CK TTP mappings for multiple threat actors.
2. A new set of detected TTPs from a recent incident.

Your task:
- Compare the input TTPs to known actor profiles.
- Identify the actor(s) with the most similar or overlapping TTPs.
- Explain your reasoning clearly and concisely.
- End with a summary and a confidence rating (High/Medium/Low).

Known Actor Mappings:
{summarized_mapping}

Detected Input TTPs:
{', '.join(input_ttps)}

Return your analysis formatted as follows:
---
**Analysis Summary:**
[Your explanation]

**Overlap Table:**
[List top 3 matching actors and number of shared TTPs]

**Most Likely Attacker:**
[Actor Name] – Confidence: [High/Medium/Low]
---
"""

    # Run OpenAI analysis
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.3,
        messages=[
            {"role": "system", "content": "You are an experienced MITRE ATT&CK threat analyst."},
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content.strip()

def generate_word_report(report_text, input_ttps):
    """Save the AI-generated analysis into a Word document."""
    doc = Document()

    # Title
    doc.add_heading("Threat Attribution Report", level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Detected TTPs: {', '.join(input_ttps)}")
    doc.add_paragraph("")

    # Content
    doc.add_heading("AI-Generated Analysis", level=2)
    doc.add_paragraph(report_text)

    # Save document
    filename = f"Threat_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    print(f"✅ Report saved as {filename}")

if __name__ == "__main__":
    # Example TTPs (replace with your input from Flask)
    pseudo_input_ttp = ["T1059", "T1078", "T1071"]
    mapping_path = "actor_ttp_mapping.json" # Path to your JSON mapping file, edit later

    print("🔍 Running Analysis...")
    report = analyze_TTP(mapping_path, pseudo_input_ttp)

    # Save as Word document
    generate_word_report(report, pseudo_input_ttp)
